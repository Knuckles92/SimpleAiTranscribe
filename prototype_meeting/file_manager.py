"""
File Manager for Meeting Transcription System

Handles persistent file storage, organization, and management for meeting recordings
and transcripts. Provides atomic file operations, backup functionality, and
export capabilities.
"""

import os
import json
import shutil
import tempfile
import logging
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass, asdict
import hashlib
import zipfile
from contextlib import contextmanager

# Import configuration
from .config_meeting import meeting_config

# Optional imports for export functionality
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


@dataclass
class FileMetadata:
    """Metadata for a file."""
    file_path: str
    file_type: str
    size_bytes: int
    created_at: datetime
    modified_at: datetime
    checksum: str
    session_id: str


@dataclass
class SessionInfo:
    """Information about a meeting session."""
    session_id: str
    timestamp: str
    directory_path: str
    audio_file: Optional[str]
    transcript_file: Optional[str]
    metadata_file: Optional[str]
    chunks_directory: Optional[str]
    created_at: datetime
    total_size_bytes: int
    file_count: int


class FileManagerError(Exception):
    """Base exception for FileManager operations."""
    pass


class FileManager:
    """
    Manages persistent file storage and organization for meeting recordings.
    
    Features:
    - Atomic file operations to prevent corruption
    - Organized directory structure by date/session
    - Backup and restore capabilities
    - Multiple export formats
    - Metadata tracking
    - Storage statistics
    - Cleanup management
    """
    
    def __init__(self, base_directory: Optional[str] = None, enable_logging: bool = True):
        """
        Initialize FileManager.
        
        Args:
            base_directory: Base directory for all meeting files
            enable_logging: Whether to enable logging
        """
        self.base_directory = Path(base_directory or meeting_config.MEETINGS_BASE_DIR)
        self.backup_directory = self.base_directory / "backups"
        self.temp_directory = self.base_directory / "temp"
        
        # Setup logging
        if enable_logging:
            self._setup_logging()
        
        # Ensure base directories exist
        self._ensure_base_directories()
    
    def _setup_logging(self):
        """Setup logging for file operations."""
        log_dir = self.base_directory / "logs"
        log_dir.mkdir(parents=True, exist_ok=True)
        
        log_file = log_dir / "file_manager.log"
        
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(log_file),
                logging.StreamHandler()
            ]
        )
        
        self.logger = logging.getLogger(__name__)
        self.logger.info("FileManager initialized")
    
    def _ensure_base_directories(self):
        """Ensure all base directories exist."""
        directories = [
            self.base_directory,
            self.backup_directory,
            self.temp_directory,
            self.base_directory / "exports"
        ]
        
        for directory in directories:
            directory.mkdir(parents=True, exist_ok=True)
    
    @contextmanager
    def _atomic_write(self, file_path: Path):
        """
        Context manager for atomic file writes.
        
        Writes to a temporary file first, then moves to the target location
        to prevent corruption if the operation is interrupted.
        """
        file_path = Path(file_path)
        temp_path = self.temp_directory / f"{file_path.name}.tmp.{os.getpid()}"
        
        try:
            with open(temp_path, 'wb') as temp_file:
                yield temp_file
            
            # Ensure parent directory exists
            file_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Atomic move
            shutil.move(str(temp_path), str(file_path))
            
        except Exception as e:
            # Clean up temp file on error
            if temp_path.exists():
                temp_path.unlink()
            raise FileManagerError(f"Atomic write failed for {file_path}: {e}")
    
    def _calculate_checksum(self, file_path: Path) -> str:
        """Calculate MD5 checksum of a file."""
        hash_md5 = hashlib.md5()
        try:
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except Exception as e:
            self.logger.error(f"Failed to calculate checksum for {file_path}: {e}")
            return ""
    
    def _get_file_metadata(self, file_path: Path, session_id: str) -> FileMetadata:
        """Get metadata for a file."""
        stat = file_path.stat()
        
        return FileMetadata(
            file_path=str(file_path),
            file_type=file_path.suffix.lower(),
            size_bytes=stat.st_size,
            created_at=datetime.fromtimestamp(stat.st_ctime),
            modified_at=datetime.fromtimestamp(stat.st_mtime),
            checksum=self._calculate_checksum(file_path),
            session_id=session_id
        )
    
    def create_meeting_directory(self, session_id: str) -> str:
        """
        Create directory structure for a meeting session.
        
        Args:
            session_id: Unique session identifier (timestamp format)
            
        Returns:
            Path to the created session directory
        """
        try:
            # Parse date from session_id for organization
            date_str = session_id.split('_')[0]  # Extract YYYYMMDD
            year = date_str[:4]
            month = date_str[4:6]
            day = date_str[6:8]
            
            # Create hierarchical directory structure
            session_dir = self.base_directory / year / month / day / f"meeting_{session_id}"
            session_dir.mkdir(parents=True, exist_ok=True)
            
            # Create subdirectories
            chunks_dir = session_dir / meeting_config.MEETING_CHUNKS_DIR
            chunks_dir.mkdir(exist_ok=True)
            
            exports_dir = session_dir / "exports"
            exports_dir.mkdir(exist_ok=True)
            
            self.logger.info(f"Created meeting directory: {session_dir}")
            return str(session_dir)
            
        except Exception as e:
            raise FileManagerError(f"Failed to create meeting directory for {session_id}: {e}")
    
    def get_session_directory(self, session_id: str) -> Path:
        """Get the directory path for a session."""
        date_str = session_id.split('_')[0]
        year = date_str[:4]
        month = date_str[4:6]
        day = date_str[6:8]
        
        return self.base_directory / year / month / day / f"meeting_{session_id}"
    
    def get_audio_file_path(self, session_id: str) -> str:
        """
        Get the path for the audio file of a session.
        
        Args:
            session_id: Session identifier
            
        Returns:
            Full path to the audio file
        """
        session_dir = self.get_session_directory(session_id)
        filename = meeting_config.MEETING_AUDIO_TEMPLATE.format(timestamp=session_id)
        return str(session_dir / filename)
    
    def get_transcript_file_path(self, session_id: str) -> str:
        """
        Get the path for the transcript file of a session.
        
        Args:
            session_id: Session identifier
            
        Returns:
            Full path to the transcript file
        """
        session_dir = self.get_session_directory(session_id)
        filename = meeting_config.MEETING_TRANSCRIPT_TEMPLATE.format(timestamp=session_id)
        return str(session_dir / filename)
    
    def get_metadata_file_path(self, session_id: str) -> str:
        """Get the path for the metadata file of a session."""
        session_dir = self.get_session_directory(session_id)
        filename = meeting_config.MEETING_METADATA_TEMPLATE.format(timestamp=session_id)
        return str(session_dir / filename)
    
    def save_audio_file(self, session_id: str, audio_data: bytes) -> str:
        """
        Save audio data to file with proper naming and atomic operations.
        
        Args:
            session_id: Session identifier
            audio_data: Raw audio data
            
        Returns:
            Path to saved audio file
        """
        try:
            # Ensure directory exists
            self.create_meeting_directory(session_id)
            
            audio_path = Path(self.get_audio_file_path(session_id))
            
            with self._atomic_write(audio_path) as f:
                f.write(audio_data)
            
            self.logger.info(f"Saved audio file: {audio_path} ({len(audio_data)} bytes)")
            return str(audio_path)
            
        except Exception as e:
            raise FileManagerError(f"Failed to save audio file for {session_id}: {e}")
    
    def save_transcript_file(self, session_id: str, text: str, metadata: Optional[Dict[str, Any]] = None) -> str:
        """
        Save transcript text with metadata.
        
        Args:
            session_id: Session identifier
            text: Transcript text
            metadata: Optional metadata dictionary
            
        Returns:
            Path to saved transcript file
        """
        try:
            # Ensure directory exists
            self.create_meeting_directory(session_id)
            
            transcript_path = Path(self.get_transcript_file_path(session_id))
            
            # Save transcript
            with self._atomic_write(transcript_path) as f:
                f.write(text.encode('utf-8'))
            
            # Save metadata if provided
            if metadata:
                metadata_path = Path(self.get_metadata_file_path(session_id))
                metadata_with_session = {
                    **metadata,
                    'session_id': session_id,
                    'transcript_path': str(transcript_path),
                    'saved_at': datetime.now().isoformat()
                }
                
                with self._atomic_write(metadata_path) as f:
                    f.write(json.dumps(metadata_with_session, indent=2).encode('utf-8'))
            
            self.logger.info(f"Saved transcript file: {transcript_path}")
            return str(transcript_path)
            
        except Exception as e:
            raise FileManagerError(f"Failed to save transcript for {session_id}: {e}")
    
    def export_transcript(self, session_id: str, format_type: str, options: Optional[Dict[str, Any]] = None) -> str:
        """
        Export transcript to various formats.
        
        Args:
            session_id: Session identifier
            format_type: Export format ('txt', 'docx', 'pdf', 'json', 'srt', 'vtt')
            options: Export options (timestamps, chunk_markers, etc.)
            
        Returns:
            Path to exported file
        """
        if format_type not in meeting_config.EXPORT_FORMATS:
            raise FileManagerError(f"Unsupported export format: {format_type}")
        
        try:
            # Load transcript and metadata
            transcript_path = Path(self.get_transcript_file_path(session_id))
            metadata_path = Path(self.get_metadata_file_path(session_id))
            
            if not transcript_path.exists():
                raise FileManagerError(f"Transcript not found for session {session_id}")
            
            with open(transcript_path, 'r', encoding='utf-8') as f:
                transcript_text = f.read()
            
            metadata = {}
            if metadata_path.exists():
                with open(metadata_path, 'r', encoding='utf-8') as f:
                    metadata = json.load(f)
            
            # Create export directory
            session_dir = self.get_session_directory(session_id)
            exports_dir = session_dir / "exports"
            exports_dir.mkdir(exist_ok=True)
            
            # Generate export filename
            export_filename = f"meeting_{session_id}.{format_type}"
            export_path = exports_dir / export_filename
            
            # Apply export options
            options = options or {}
            include_timestamps = options.get('include_timestamps', meeting_config.INCLUDE_TIMESTAMPS)
            include_chunk_markers = options.get('include_chunk_markers', meeting_config.INCLUDE_CHUNK_MARKERS)
            
            # Export based on format
            if format_type == 'txt':
                self._export_txt(export_path, transcript_text, metadata, include_timestamps, include_chunk_markers)
            elif format_type == 'json':
                self._export_json(export_path, transcript_text, metadata)
            elif format_type == 'srt':
                self._export_srt(export_path, transcript_text, metadata)
            elif format_type == 'vtt':
                self._export_vtt(export_path, transcript_text, metadata)
            elif format_type == 'docx':
                self._export_docx(export_path, transcript_text, metadata, include_timestamps)
            elif format_type == 'pdf':
                self._export_pdf(export_path, transcript_text, metadata, include_timestamps)
            
            self.logger.info(f"Exported transcript to {format_type}: {export_path}")
            return str(export_path)
            
        except Exception as e:
            raise FileManagerError(f"Failed to export transcript for {session_id}: {e}")
    
    def _export_txt(self, export_path: Path, text: str, metadata: Dict, include_timestamps: bool, include_chunk_markers: bool):
        """Export to plain text format."""
        content = []
        
        # Add header
        if metadata.get('saved_at'):
            content.append(f"Meeting Transcript - {metadata['saved_at']}")
            content.append("=" * 50)
            content.append("")
        
        content.append(text)
        
        with self._atomic_write(export_path) as f:
            f.write('\n'.join(content).encode('utf-8'))
    
    def _export_json(self, export_path: Path, text: str, metadata: Dict):
        """Export to JSON format with full metadata."""
        json_data = {
            'transcript': text,
            'metadata': metadata,
            'exported_at': datetime.now().isoformat(),
            'format_version': '1.0'
        }
        
        with self._atomic_write(export_path) as f:
            f.write(json.dumps(json_data, indent=2).encode('utf-8'))
    
    def _export_srt(self, export_path: Path, text: str, metadata: Dict):
        """Export to SRT subtitle format."""
        # Basic SRT export - would need timestamp data for proper implementation
        content = [
            "1",
            "00:00:00,000 --> 00:00:10,000",
            text[:100] + "..." if len(text) > 100 else text,
            ""
        ]
        
        with self._atomic_write(export_path) as f:
            f.write('\n'.join(content).encode('utf-8'))
    
    def _export_vtt(self, export_path: Path, text: str, metadata: Dict):
        """Export to WebVTT format."""
        content = [
            "WEBVTT",
            "",
            "00:00:00.000 --> 00:00:10.000",
            text[:100] + "..." if len(text) > 100 else text,
            ""
        ]
        
        with self._atomic_write(export_path) as f:
            f.write('\n'.join(content).encode('utf-8'))
    
    def _export_docx(self, export_path: Path, text: str, metadata: Dict, include_timestamps: bool):
        """Export to Word document format."""
        if not DOCX_AVAILABLE:
            raise FileManagerError("python-docx package not available for DOCX export")
        
        doc = Document()
        
        # Add title
        title = doc.add_heading('Meeting Transcript', 0)
        
        # Add metadata
        if metadata.get('saved_at'):
            doc.add_paragraph(f"Date: {metadata['saved_at']}")
        
        # Add transcript
        doc.add_paragraph(text)
        
        # Save to temporary file then move
        temp_path = self.temp_directory / f"{export_path.name}.tmp"
        doc.save(str(temp_path))
        shutil.move(str(temp_path), str(export_path))
    
    def _export_pdf(self, export_path: Path, text: str, metadata: Dict, include_timestamps: bool):
        """Export to PDF format."""
        if not PDF_AVAILABLE:
            raise FileManagerError("reportlab package not available for PDF export")
        
        temp_path = self.temp_directory / f"{export_path.name}.tmp"
        
        doc = SimpleDocTemplate(str(temp_path), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Add title
        story.append(Paragraph("Meeting Transcript", styles['Title']))
        story.append(Spacer(1, 12))
        
        # Add metadata
        if metadata.get('saved_at'):
            story.append(Paragraph(f"Date: {metadata['saved_at']}", styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Add transcript
        story.append(Paragraph(text, styles['Normal']))
        
        doc.build(story)
        shutil.move(str(temp_path), str(export_path))
    
    def get_meeting_files(self, session_id: str) -> Dict[str, List[FileMetadata]]:
        """
        Get all files for a meeting session.
        
        Args:
            session_id: Session identifier
            
        Returns:
            Dictionary categorizing files by type
        """
        session_dir = self.get_session_directory(session_id)
        
        if not session_dir.exists():
            return {}
        
        files = {
            'audio': [],
            'transcripts': [],
            'metadata': [],
            'chunks': [],
            'exports': [],
            'other': []
        }
        
        try:
            for file_path in session_dir.rglob('*'):
                if file_path.is_file():
                    metadata = self._get_file_metadata(file_path, session_id)
                    
                    # Categorize file
                    if file_path.suffix.lower() in ['.wav', '.mp3', '.m4a', '.ogg']:
                        files['audio'].append(metadata)
                    elif file_path.suffix.lower() == '.txt' and 'export' not in file_path.parts:
                        files['transcripts'].append(metadata)
                    elif file_path.suffix.lower() == '.json':
                        files['metadata'].append(metadata)
                    elif 'chunk' in file_path.name.lower():
                        files['chunks'].append(metadata)
                    elif 'export' in file_path.parts:
                        files['exports'].append(metadata)
                    else:
                        files['other'].append(metadata)
        
        except Exception as e:
            self.logger.error(f"Error getting files for session {session_id}: {e}")
        
        return files
    
    def cleanup_old_meetings(self, days_old: int) -> Dict[str, int]:
        """
        Clean up meetings older than specified days.
        
        Args:
            days_old: Remove meetings older than this many days
            
        Returns:
            Dictionary with cleanup statistics
        """
        cutoff_date = datetime.now() - timedelta(days=days_old)
        stats = {'directories_removed': 0, 'files_removed': 0, 'bytes_freed': 0}
        
        try:
            for year_dir in self.base_directory.iterdir():
                if not year_dir.is_dir() or not year_dir.name.isdigit():
                    continue
                
                for month_dir in year_dir.iterdir():
                    if not month_dir.is_dir():
                        continue
                    
                    for day_dir in month_dir.iterdir():
                        if not day_dir.is_dir():
                            continue
                        
                        for session_dir in day_dir.iterdir():
                            if not session_dir.is_dir() or not session_dir.name.startswith('meeting_'):
                                continue
                            
                            # Check if session is old enough
                            session_created = datetime.fromtimestamp(session_dir.stat().st_ctime)
                            
                            if session_created < cutoff_date:
                                # Calculate size before deletion
                                session_size = sum(f.stat().st_size for f in session_dir.rglob('*') if f.is_file())
                                file_count = len(list(session_dir.rglob('*')))
                                
                                # Remove session directory
                                shutil.rmtree(session_dir)
                                
                                stats['directories_removed'] += 1
                                stats['files_removed'] += file_count
                                stats['bytes_freed'] += session_size
                                
                                self.logger.info(f"Removed old session: {session_dir}")
        
        except Exception as e:
            self.logger.error(f"Error during cleanup: {e}")
            raise FileManagerError(f"Cleanup failed: {e}")
        
        self.logger.info(f"Cleanup completed: {stats}")
        return stats
    
    def get_storage_statistics(self) -> Dict[str, Any]:
        """
        Get storage usage statistics.
        
        Returns:
            Dictionary with storage statistics
        """
        stats = {
            'total_sessions': 0,
            'total_files': 0,
            'total_size_bytes': 0,
            'by_file_type': {},
            'by_date': {},
            'oldest_session': None,
            'newest_session': None
        }
        
        try:
            session_dates = []
            
            for file_path in self.base_directory.rglob('*'):
                if file_path.is_file():
                    file_size = file_path.stat().st_size
                    file_ext = file_path.suffix.lower()
                    
                    stats['total_files'] += 1
                    stats['total_size_bytes'] += file_size
                    
                    # Track by file type
                    if file_ext not in stats['by_file_type']:
                        stats['by_file_type'][file_ext] = {'count': 0, 'size_bytes': 0}
                    stats['by_file_type'][file_ext]['count'] += 1
                    stats['by_file_type'][file_ext]['size_bytes'] += file_size
                    
                    # Track session info
                    if 'meeting_' in str(file_path):
                        # Extract session date from path
                        parts = str(file_path).split(os.sep)
                        for part in parts:
                            if part.startswith('meeting_') and '_' in part:
                                try:
                                    timestamp = part.replace('meeting_', '').split('_')[0]
                                    session_date = datetime.strptime(timestamp, '%Y%m%d')
                                    session_dates.append(session_date)
                                    break
                                except ValueError:
                                    continue
            
            # Count unique sessions
            for session_dir in self.base_directory.rglob('meeting_*'):
                if session_dir.is_dir():
                    stats['total_sessions'] += 1
            
            # Date statistics
            if session_dates:
                stats['oldest_session'] = min(session_dates).isoformat()
                stats['newest_session'] = max(session_dates).isoformat()
        
        except Exception as e:
            self.logger.error(f"Error calculating storage statistics: {e}")
        
        return stats
    
    def backup_session(self, session_id: str) -> str:
        """
        Create a backup of a session.
        
        Args:
            session_id: Session identifier
            
        Returns:
            Path to backup file
        """
        session_dir = self.get_session_directory(session_id)
        
        if not session_dir.exists():
            raise FileManagerError(f"Session directory not found: {session_id}")
        
        try:
            # Create backup filename with timestamp
            backup_timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            backup_filename = f"backup_{session_id}_{backup_timestamp}.zip"
            backup_path = self.backup_directory / backup_filename
            
            # Create zip backup
            with zipfile.ZipFile(backup_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for file_path in session_dir.rglob('*'):
                    if file_path.is_file():
                        arcname = file_path.relative_to(session_dir)
                        zipf.write(file_path, arcname)
            
            self.logger.info(f"Created backup: {backup_path}")
            return str(backup_path)
        
        except Exception as e:
            raise FileManagerError(f"Failed to backup session {session_id}: {e}")
    
    def restore_session_from_backup(self, backup_path: str) -> str:
        """
        Restore a session from backup.
        
        Args:
            backup_path: Path to backup zip file
            
        Returns:
            Path to restored session directory
        """
        backup_path = Path(backup_path)
        
        if not backup_path.exists():
            raise FileManagerError(f"Backup file not found: {backup_path}")
        
        try:
            # Extract session ID from backup filename
            backup_name = backup_path.stem
            parts = backup_name.split('_')
            if len(parts) < 3:
                raise FileManagerError("Invalid backup filename format")
            
            # Reconstruct session_id (everything between 'backup_' and the last timestamp)
            session_id = '_'.join(parts[1:-2])  # Skip 'backup' and timestamp
            
            # Create restoration directory
            restored_dir = self.get_session_directory(session_id)
            
            # Check if session already exists
            if restored_dir.exists():
                # Create a backup of existing session
                existing_backup = self.backup_session(session_id)
                self.logger.info(f"Backed up existing session before restore: {existing_backup}")
            
            # Extract backup
            with zipfile.ZipFile(backup_path, 'r') as zipf:
                zipf.extractall(restored_dir)
            
            self.logger.info(f"Restored session from backup: {restored_dir}")
            return str(restored_dir)
        
        except Exception as e:
            raise FileManagerError(f"Failed to restore from backup {backup_path}: {e}")
    
    def get_session_info(self, session_id: str) -> Optional[SessionInfo]:
        """
        Get comprehensive information about a session.
        
        Args:
            session_id: Session identifier
            
        Returns:
            SessionInfo object or None if session doesn't exist
        """
        session_dir = self.get_session_directory(session_id)
        
        if not session_dir.exists():
            return None
        
        try:
            # Collect file information
            audio_file = None
            transcript_file = None
            metadata_file = None
            chunks_directory = None
            total_size = 0
            file_count = 0
            
            audio_path = Path(self.get_audio_file_path(session_id))
            if audio_path.exists():
                audio_file = str(audio_path)
            
            transcript_path = Path(self.get_transcript_file_path(session_id))
            if transcript_path.exists():
                transcript_file = str(transcript_path)
            
            metadata_path = Path(self.get_metadata_file_path(session_id))
            if metadata_path.exists():
                metadata_file = str(metadata_path)
            
            chunks_dir = session_dir / meeting_config.MEETING_CHUNKS_DIR
            if chunks_dir.exists():
                chunks_directory = str(chunks_dir)
            
            # Calculate total size and file count
            for file_path in session_dir.rglob('*'):
                if file_path.is_file():
                    total_size += file_path.stat().st_size
                    file_count += 1
            
            # Get creation timestamp
            created_at = datetime.fromtimestamp(session_dir.stat().st_ctime)
            
            return SessionInfo(
                session_id=session_id,
                timestamp=session_id,
                directory_path=str(session_dir),
                audio_file=audio_file,
                transcript_file=transcript_file,
                metadata_file=metadata_file,
                chunks_directory=chunks_directory,
                created_at=created_at,
                total_size_bytes=total_size,
                file_count=file_count
            )
        
        except Exception as e:
            self.logger.error(f"Error getting session info for {session_id}: {e}")
            return None
    
    def list_all_sessions(self) -> List[SessionInfo]:
        """
        List all meeting sessions.
        
        Returns:
            List of SessionInfo objects
        """
        sessions = []
        
        try:
            for session_dir in self.base_directory.rglob('meeting_*'):
                if session_dir.is_dir():
                    # Extract session ID from directory name
                    session_id = session_dir.name.replace('meeting_', '')
                    session_info = self.get_session_info(session_id)
                    if session_info:
                        sessions.append(session_info)
        
        except Exception as e:
            self.logger.error(f"Error listing sessions: {e}")
        
        # Sort by creation date (newest first)
        sessions.sort(key=lambda x: x.created_at, reverse=True)
        return sessions
    
    def cleanup_temp_files(self):
        """Clean up temporary files."""
        try:
            if self.temp_directory.exists():
                for temp_file in self.temp_directory.iterdir():
                    if temp_file.is_file():
                        # Remove files older than 1 hour
                        if datetime.now().timestamp() - temp_file.stat().st_mtime > 3600:
                            temp_file.unlink()
                            self.logger.info(f"Removed old temp file: {temp_file}")
        
        except Exception as e:
            self.logger.error(f"Error cleaning up temp files: {e}")


# Convenience function to create a default FileManager instance
def create_file_manager(base_directory: Optional[str] = None) -> FileManager:
    """Create a FileManager instance with default settings."""
    return FileManager(base_directory)


if __name__ == "__main__":
    # Example usage
    fm = create_file_manager()
    
    # Create a test session
    session_id = meeting_config.get_meeting_timestamp()
    print(f"Created session: {session_id}")
    
    # Create directory
    session_dir = fm.create_meeting_directory(session_id)
    print(f"Session directory: {session_dir}")
    
    # Get file paths
    audio_path = fm.get_audio_file_path(session_id)
    transcript_path = fm.get_transcript_file_path(session_id)
    print(f"Audio path: {audio_path}")
    print(f"Transcript path: {transcript_path}")
    
    # Get storage statistics
    stats = fm.get_storage_statistics()
    print(f"Storage statistics: {stats}")